# ==============================================================================
# format_thesis.py — 论文格式转换工具（完整版）
# 功能：图形界面 / 智能标题识别 / 字体字号 / 首行缩进 / 段落间距 / 格式检查 / 多套模板
# 依赖：pip install python-docx
# 运行：python format_thesis.py
# ==============================================================================

import re
import json
import os
import copy
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==============================================================================
# 默认模板（首次运行自动写入 config.json，之后从文件读取）
# ==============================================================================

DEFAULT_CONFIG = {
    "templates": {
        "通用模板": {
            "chinese_font": "宋体",
            "western_font": "Times New Roman",
            "sizes": {
                "main_title": 32,
                "heading1":   30,
                "heading2":   28,
                "heading3":   24,
                "body":       21
            },
            "line_spacing": 1.5,
            "spacing": {
                "main_title": [24, 12],
                "heading1":   [24, 6],
                "heading2":   [18, 6],
                "heading3":   [12, 6],
                "body":       [0,  0]
            },
            "first_line_indent": True
        },
        "学术期刊投稿": {
            "chinese_font": "宋体",
            "western_font": "Times New Roman",
            "sizes": {
                "main_title": 32,
                "heading1":   28,
                "heading2":   26,
                "heading3":   24,
                "body":       24
            },
            "line_spacing": 2.0,
            "spacing": {
                "main_title": [12, 12],
                "heading1":   [12, 6],
                "heading2":   [6,  6],
                "heading3":   [6,  3],
                "body":       [0,  0]
            },
            "first_line_indent": True
        }
    }
}

CONFIG_FILE = "config.json"


# ==============================================================================
# 配置读写
# ==============================================================================

def load_config() -> dict:
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    save_config(DEFAULT_CONFIG)
    return DEFAULT_CONFIG


def save_config(config: dict):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


# ==============================================================================
# 标题识别（基于编号规则）
# ==============================================================================

RE_HEADING3 = re.compile(r"^[0-9０-９]+[\.。][0-9０-９]+[\.。][0-9０-９]+")
RE_HEADING2 = re.compile(r"^[0-9０-９]+[\.。][0-9０-９]+(?![\.。][0-9０-９])")
RE_HEADING1 = re.compile(
    r"^("
    r"[0-9０-９]+[\s　\.。、]"
    r"|[一二三四五六七八九十百]+[、\.]"
    r"|第[一二三四五六七八九十百0-9]+[章节部篇]"
    r")"
)

HEADING_MAX_LEN    = 40
MAIN_TITLE_MAX_LEN = 30


def get_run_size(run) -> int:
    """读取 run 字号（半磅），读不到返回 0"""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is not None:
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            return int(sz.get(qn("w:val"), 0))
    return 0


def para_max_size(para) -> int:
    sizes = [get_run_size(r) for r in para.runs]
    return max(sizes) if sizes else 0


def para_is_bold(para) -> bool:
    for run in para.runs:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None:
            b = rPr.find(qn("w:b"))
            if b is not None and b.get(qn("w:val"), "1") != "0":
                return True
    return False


def detect_level(para, all_sizes: list, idx: int) -> str:
    """
    按优先级判断段落层级：
      1. 超长段落 → 正文
      2. 匹配三级编号 X.X.X → heading3
      3. 匹配二级编号 X.X   → heading2
      4. 匹配一级编号 X     → heading1
      5. 短段落且突出（加粗/字号最大）→ main_title
      6. 其他 → body
    """
    text = para.text.strip()
    if not text or len(text) > HEADING_MAX_LEN:
        return "body"
    if RE_HEADING3.match(text):
        return "heading3"
    if RE_HEADING2.match(text):
        return "heading2"
    if RE_HEADING1.match(text):
        return "heading1"
    if len(text) <= MAIN_TITLE_MAX_LEN:
        my_size   = para_max_size(para)
        prev_size = all_sizes[idx - 1] if idx > 0 else 0
        next_size = all_sizes[idx + 1] if idx < len(all_sizes) - 1 else 0
        if para_is_bold(para) or (my_size > 0 and my_size > prev_size and my_size > next_size):
            return "main_title"
    return "body"


# ==============================================================================
# 格式应用
# ==============================================================================

def apply_run_format(run, size_half_pt: int, chinese_font: str, western_font: str):
    """
    写入字体和字号。
    w:rFonts 同时设中西文字体，Word 按 Unicode 范围自动分配，
    不需要手动拆分中英文 run。
    """
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._r.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(qn("w:ascii"),    western_font)
    rFonts.set(qn("w:hAnsi"),    western_font)
    rFonts.set(qn("w:eastAsia"), chinese_font)
    rFonts.set(qn("w:cs"),       western_font)

    for tag in (qn("w:sz"), qn("w:szCs")):
        elem = rPr.find(tag)
        if elem is None:
            elem = OxmlElement(tag)
            rPr.append(elem)
        elem.set(qn("w:val"), str(size_half_pt))


def apply_para_format(para, level: str, template: dict):
    """
    写入段落格式：行距、段前段后间距、首行缩进。
    - 标题不加首行缩进
    - 正文按 line_spacing 设定行距，标题固定单倍
    - 间距单位 pt，内部转 twips（1pt = 20 twips）
    """
    size_half_pt = template["sizes"][level]
    spacing_cfg  = template["spacing"]
    line_spacing = template["line_spacing"]
    do_indent    = template["first_line_indent"]

    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._p.insert(0, pPr)

    # ── 段前段后 + 行距 ───────────────────────────────────────────────────────
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)

    before_pt, after_pt = spacing_cfg.get(level, [0, 0])
    sp.set(qn("w:before"), str(int(before_pt * 20)))
    sp.set(qn("w:after"),  str(int(after_pt  * 20)))

    # 正文用配置行距，标题统一单倍行距
    multiplier = line_spacing if level == "body" else 1.0
    sp.set(qn("w:line"),     str(int(240 * multiplier)))
    sp.set(qn("w:lineRule"), "auto")

    # ── 首行缩进 ─────────────────────────────────────────────────────────────
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)

    if do_indent and level == "body":
        # 2个字符宽度 = 2 × 字号pt × 20 twips
        # size_half_pt / 2 = 字号pt，× 2字符 × 20 twips/pt = size_half_pt * 20
        indent_twips = size_half_pt * 20
        ind.set(qn("w:firstLine"), str(indent_twips))
    else:
        ind.attrib.pop(qn("w:firstLine"),      None)
        ind.attrib.pop(qn("w:firstLineChars"), None)


# ==============================================================================
# 格式检查（转换前生成报告）
# ==============================================================================

def check_format(doc_path: str, template: dict) -> list:
    """
    扫描文档，找出与目标模板不符的段落。
    返回列表，每项: {"level": ..., "text": ..., "issues": [...]}
    """
    doc = Document(doc_path)
    paragraphs = doc.paragraphs
    all_sizes  = [para_max_size(p) for p in paragraphs]
    target_sizes = template["sizes"]
    issues = []

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text or not para.runs:
            continue

        level = detect_level(para, all_sizes, i)
        target_size = target_sizes[level]
        para_issues = []

        # 检查字号
        for run in para.runs:
            current = get_run_size(run)
            if current != 0 and current != target_size:
                para_issues.append(
                    f"字号应为 {target_size / 2}pt，当前为 {current / 2}pt"
                )
                break

        # 检查首行缩进（仅正文）
        if level == "body" and template["first_line_indent"]:
            pPr = para._p.find(qn("w:pPr"))
            has_indent = False
            if pPr is not None:
                ind = pPr.find(qn("w:ind"))
                if ind is not None and ind.get(qn("w:firstLine")):
                    has_indent = True
            if not has_indent:
                para_issues.append("正文缺少首行缩进")

        if para_issues:
            issues.append({
                "level": level,
                "text":  text[:40],
                "issues": para_issues
            })

    return issues


# ==============================================================================
# 主转换流程
# ==============================================================================

LEVEL_NAMES = {
    "main_title": "主标题",
    "heading1":   "一级标题",
    "heading2":   "二级标题",
    "heading3":   "三级标题",
    "body":       "正文",
}


def convert_document(input_path: str, output_path: str, template: dict, log_fn):
    """
    执行完整转换：识别层级 → 写字体字号 → 写段落格式 → 保存。
    log_fn 是回调，用于向 GUI 日志区输出信息。
    """
    doc = Document(input_path)
    paragraphs = doc.paragraphs
    all_sizes   = [para_max_size(p) for p in paragraphs]

    chinese_font = template["chinese_font"]
    western_font = template["western_font"]
    sizes        = template["sizes"]

    count = 0
    for i, para in enumerate(paragraphs):
        if not para.text.strip() or not para.runs:
            continue

        level = detect_level(para, all_sizes, i)

        for run in para.runs:
            apply_run_format(run, sizes[level], chinese_font, western_font)

        apply_para_format(para, level, template)

        count += 1
        preview = para.text[:42] + ("..." if len(para.text) > 42 else "")
        log_fn(f"  [{LEVEL_NAMES[level]:5}] {preview}")

    doc.save(output_path)
    log_fn(f"\n✅ 完成！共处理 {count} 个段落 → {output_path}")


# ==============================================================================
# 模板编辑弹窗
# ==============================================================================

class TemplateEditor(tk.Toplevel):
    """
    弹出窗口，用于新建或编辑模板。
    编辑完点保存后，通过 callback 把新模板名和内容传回主窗口。
    """
    SIZE_FIELDS = [
        ("主标题字号（半磅，16pt=32）", "main_title"),
        ("一级标题字号（15pt=30）",     "heading1"),
        ("二级标题字号（14pt=28）",     "heading2"),
        ("三级标题字号（12pt=24）",     "heading3"),
        ("正文字号（10.5pt=21）",       "body"),
    ]

    def __init__(self, parent, name: str, template: dict, callback, is_new=False):
        super().__init__(parent)
        self.title("新建模板" if is_new else f"编辑模板：{name}")
        self.resizable(False, False)
        self.grab_set()

        self.template = copy.deepcopy(template)
        self.callback = callback
        self.is_new   = is_new

        row = 0
        pad = {"padx": 14, "pady": 5}

        # 模板名称
        ttk.Label(self, text="模板名称").grid(row=row, column=0, sticky="w", **pad)
        self.name_var = tk.StringVar(value="新模板" if is_new else name)
        ttk.Entry(self, textvariable=self.name_var, width=24).grid(row=row, column=1, **pad)
        row += 1

        # 字号字段
        self.size_vars = {}
        for label, key in self.SIZE_FIELDS:
            ttk.Label(self, text=label).grid(row=row, column=0, sticky="w", **pad)
            var = tk.StringVar(value=str(template["sizes"][key]))
            ttk.Entry(self, textvariable=var, width=24).grid(row=row, column=1, **pad)
            self.size_vars[key] = var
            row += 1

        # 字体
        ttk.Label(self, text="中文字体").grid(row=row, column=0, sticky="w", **pad)
        self.cn_var = tk.StringVar(value=template["chinese_font"])
        ttk.Entry(self, textvariable=self.cn_var, width=24).grid(row=row, column=1, **pad)
        row += 1

        ttk.Label(self, text="西文字体").grid(row=row, column=0, sticky="w", **pad)
        self.en_var = tk.StringVar(value=template["western_font"])
        ttk.Entry(self, textvariable=self.en_var, width=24).grid(row=row, column=1, **pad)
        row += 1

        # 行距
        ttk.Label(self, text="正文行距倍数（如 1.5）").grid(row=row, column=0, sticky="w", **pad)
        self.ls_var = tk.StringVar(value=str(template["line_spacing"]))
        ttk.Entry(self, textvariable=self.ls_var, width=24).grid(row=row, column=1, **pad)
        row += 1

        # 首行缩进
        self.indent_var = tk.BooleanVar(value=template["first_line_indent"])
        ttk.Checkbutton(self, text="正文首行缩进两字符",
                        variable=self.indent_var).grid(
            row=row, column=0, columnspan=2, sticky="w", **pad)
        row += 1

        ttk.Button(self, text="保存", command=self._save, width=16).grid(
            row=row, column=0, columnspan=2, pady=12)

    def _save(self):
        name = self.name_var.get().strip()
        if not name:
            messagebox.showerror("错误", "模板名称不能为空", parent=self)
            return

        tmpl = copy.deepcopy(self.template)
        tmpl["chinese_font"]      = self.cn_var.get().strip()
        tmpl["western_font"]      = self.en_var.get().strip()
        tmpl["first_line_indent"] = self.indent_var.get()

        try:
            tmpl["line_spacing"] = float(self.ls_var.get())
        except ValueError:
            messagebox.showerror("错误", "行距请填数字，如 1.5", parent=self)
            return

        for key, var in self.size_vars.items():
            try:
                tmpl["sizes"][key] = int(var.get())
            except ValueError:
                messagebox.showerror("错误", f"字号请填整数（半磅）", parent=self)
                return

        self.callback(name, tmpl)
        self.destroy()


# ==============================================================================
# 主窗口
# ==============================================================================

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("论文格式转换工具")
        self.resizable(False, False)
        self.config_data = load_config()
        self._build_ui()

    def _build_ui(self):
        pad = {"padx": 10, "pady": 6}

        # ── 文件选择 ───────────────────────────────────────────────────────────
        f = ttk.LabelFrame(self, text="文件", padding=10)
        f.grid(row=0, column=0, sticky="ew", **pad)

        ttk.Label(f, text="输入文件:").grid(row=0, column=0, sticky="w")
        self.input_var = tk.StringVar()
        ttk.Entry(f, textvariable=self.input_var, width=48).grid(
            row=0, column=1, padx=6)
        ttk.Button(f, text="选择...",
                   command=self._pick_input).grid(row=0, column=2)

        ttk.Label(f, text="输出文件:").grid(
            row=1, column=0, sticky="w", pady=(6, 0))
        self.output_var = tk.StringVar()
        ttk.Entry(f, textvariable=self.output_var, width=48).grid(
            row=1, column=1, padx=6, pady=(6, 0))
        ttk.Button(f, text="选择...",
                   command=self._pick_output).grid(row=1, column=2, pady=(6, 0))

        # ── 模板选择 ───────────────────────────────────────────────────────────
        t = ttk.LabelFrame(self, text="格式模板", padding=10)
        t.grid(row=1, column=0, sticky="ew", **pad)

        ttk.Label(t, text="当前模板:").grid(row=0, column=0, sticky="w")
        self.tmpl_var = tk.StringVar()
        self.tmpl_combo = ttk.Combobox(
            t, textvariable=self.tmpl_var,
            values=list(self.config_data["templates"].keys()),
            state="readonly", width=22
        )
        self.tmpl_combo.grid(row=0, column=1, padx=6)
        self.tmpl_combo.current(0)

        ttk.Button(t, text="编辑模板",
                   command=self._edit_template).grid(row=0, column=2, padx=4)
        ttk.Button(t, text="新建模板",
                   command=self._new_template).grid(row=0, column=3, padx=4)

        # ── 操作按钮 ───────────────────────────────────────────────────────────
        b = ttk.Frame(self)
        b.grid(row=2, column=0, pady=6)

        ttk.Button(b, text="📋  格式检查", width=16,
                   command=self._check).pack(side="left", padx=8)
        ttk.Button(b, text="✅  开始转换", width=16,
                   command=self._convert).pack(side="left", padx=8)
        ttk.Button(b, text="清空日志", width=10,
                   command=self._clear_log).pack(side="left", padx=8)

        # ── 日志区 ─────────────────────────────────────────────────────────────
        lg = ttk.LabelFrame(self, text="运行日志", padding=10)
        lg.grid(row=3, column=0, sticky="nsew", **pad)

        self.log = scrolledtext.ScrolledText(
            lg, width=72, height=22,
            font=("Consolas", 10), state="disabled"
        )
        self.log.pack()

    # ── 文件选择回调 ──────────────────────────────────────────────────────────

    def _pick_input(self):
        path = filedialog.askopenfilename(
            title="选择输入文件",
            filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            self.input_var.set(path)
            base, ext = os.path.splitext(path)
            self.output_var.set(base + "_已格式化" + ext)

    def _pick_output(self):
        path = filedialog.asksaveasfilename(
            title="保存输出文件",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")]
        )
        if path:
            self.output_var.set(path)

    # ── 格式检查 ──────────────────────────────────────────────────────────────

    def _check(self):
        inp = self.input_var.get().strip()
        if not inp:
            messagebox.showwarning("提示", "请先选择输入文件")
            return
        if not os.path.exists(inp):
            messagebox.showerror("错误", "输入文件不存在")
            return

        tmpl = self._get_template()
        self._log("=" * 56)
        self._log("【格式检查报告】")
        self._log("=" * 56)

        try:
            issues = check_format(inp, tmpl)
        except Exception as e:
            self._log(f"检查出错：{e}")
            return

        if not issues:
            self._log("✅ 未发现问题，文档已符合目标模板要求。\n")
        else:
            self._log(f"发现 {len(issues)} 处需要调整：\n")
            for item in issues:
                self._log(f"  [{LEVEL_NAMES[item['level']]}] 「{item['text']}」")
                for iss in item["issues"]:
                    self._log(f"    → {iss}")
            self._log("")

    # ── 开始转换 ──────────────────────────────────────────────────────────────

    def _convert(self):
        inp = self.input_var.get().strip()
        out = self.output_var.get().strip()

        if not inp or not out:
            messagebox.showwarning("提示", "请填写输入和输出文件路径")
            return
        if not os.path.exists(inp):
            messagebox.showerror("错误", "输入文件不存在")
            return

        tmpl = self._get_template()
        self._log("=" * 56)
        self._log(f"【开始转换】使用模板：{self.tmpl_var.get()}")
        self._log("=" * 56)

        try:
            convert_document(inp, out, tmpl, self._log)
        except Exception as e:
            self._log(f"转换出错：{e}")

    # ── 模板管理 ──────────────────────────────────────────────────────────────

    def _get_template(self) -> dict:
        return self.config_data["templates"][self.tmpl_var.get()]

    def _edit_template(self):
        name = self.tmpl_var.get()
        TemplateEditor(self, name,
                       self.config_data["templates"][name],
                       self._on_template_saved)

    def _new_template(self):
        base = copy.deepcopy(list(self.config_data["templates"].values())[0])
        TemplateEditor(self, "", base, self._on_template_saved, is_new=True)

    def _on_template_saved(self, name: str, tmpl: dict):
        self.config_data["templates"][name] = tmpl
        save_config(self.config_data)
        self.tmpl_combo["values"] = list(self.config_data["templates"].keys())
        self.tmpl_var.set(name)
        self._log(f"✅ 模板「{name}」已保存到 config.json\n")

    # ── 日志 ──────────────────────────────────────────────────────────────────

    def _log(self, msg: str):
        self.log.configure(state="normal")
        self.log.insert("end", msg + "\n")
        self.log.see("end")
        self.log.configure(state="disabled")
        self.update_idletasks()

    def _clear_log(self):
        self.log.configure(state="normal")
        self.log.delete("1.0", "end")
        self.log.configure(state="disabled")


# ==============================================================================
# 程序入口
# ==============================================================================

if __name__ == "__main__":
    app = App()
    app.mainloop()